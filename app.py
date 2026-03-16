from __future__ import annotations

import io
import re
import html
from pathlib import Path
from typing import List, Tuple, Dict, Any, Set

import pandas as pd
import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from transformers import pipeline
import torch


# ============================================================
# DEFAULT CONFIG
# ============================================================
REDACT_TOKEN = "[redacted]"

DEFAULT_DATA_SUBJECT_FULL = "Jonathan Price"
DEFAULT_DATA_SUBJECT_EMAIL = "jonathan.price@emailclient.com"
DEFAULT_ALIASES = [
    "Jonathan Price",
    "Mr Price",
    "Mr. Price",
]

DEFAULT_KEEP_EMAILS = {
    "parking.support@heathrow.com",
    "dsar@heathrow.com",
    "dsar.team@heathrow.com",
    "legal.support@heathrow.com",
    "hr.sharedservices@heathrow.com",
    "l&d.team@heathrow.com",
    "panel.calibration@heathrow.com",
}

DEFAULT_PRIVATE_PATTERNS = [
    r"\bmy partner was at work\b",
    r"\bmy partner'?s?\b",
    r"\bGP slip is for\s+\d{1,2}\s+[A-Za-z]{3,9}\b",
    r"\bGP slip\b",
    r"\band a GP slip for your son\??\b",
    r"\bthird-party PD\b",
    r"\b\[Redacted\s*[–—-]\s*third-party PD\]\b",
    r"\bchildcare\b",
    r"\bmy son had a sudden fever\b",
    r"\bmy phone died\b",
    r"\bwith a colleague\b",
    r"\bwith union rep(?:resentative)?\b",
    r"\bwith a union representative\b",
    r"\bunion rep(?:resentative)?\b",
    r"\battended with a union representative\b",
    r"\battended with a colleague\b",
    r"\"?I asked the passenger to remove liquids.*?voice\.?\"?",
    r"\"?During secondary check,.*?offered help\.?\"?",
    r"\"?I observed from a short distance.*?speak to the passenger\.?\"?",
    r"\"?I didn'?t hear voices.*?arranging items\.?\"?",
    r"\"?Team, high load.*?escalations asap\.?\"?",
    r"\"?I kept the line moving.*?in under a minute\.?\"?",
    r"\"?I opened the toiletry bag.*?repacked immediately\.?\"?",
    r"\"?CR logged formal complaint.*?statements\.?\"?",
    r"\"?Please provide brief statements.*?apology given\.?\"?",
    r"\"?CCTV shows brief exchange.*?signage in place\.?\"?",
    r"\"?Lane 3 is live.*?courteous please\.?\"?",
    r"\"?Passenger unhappy about tone.*?repack bench\.?\"?",
    r"\"?Noted\.?\s*Keep calm.*?repacking\.?\"?",
    r"\"?Repacked toiletries;?\s*apologised\.?\"?",
    r"\"?Log in daily notes.*?check CCTV\.?\"?",
    r"\bPeak\s+\d{2}:\d{2}[–-]\d{2}:\d{2}\s+manageable\b",
    r"\bAction:\s*Review CCTV.*?phrasing\.?\b",
    r"\bCheckpoint Lead:.*?\(\d{2}:\d{2}[–-]\d{2}:\d{2}\)",
    r"\bOfficer #\d:.*?\(\d{2}:\d{2}[–-]\d{2}:\d{2}\)",
    r"\bQueue Marshal:.*?\(\d{2}:\d{2}[–-]\d{2}:\d{2}\)",
    r"\bKhan, Collins, Grant badged into.*?anomalies\.\)",
    r"\bCoaching for Officers?\s+\w+\s+and\s+\w+\s+on\b.*?handling\.?",
    r"\bSenior Screener\s*\(Lane\s*\d\)\s*present at.*?\(\d{1,2}\s+\w+\)",
    r"\bCheckpoint Lead\s*\(oversaw rota allocations\).*?\(\d{1,2}\s+\w+\)",
    r"\bNew Starter\s*\(coached by \w+\).*?\(\d{1,2}\s+\w+\)",
    r"\bUnion Rep\s*observing one feedback meeting.*?\(\d{1,2}\s+\w+\)",
    r"\bcoached by [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b",
    r"\bpresent at multiple briefings\b",
    r"\bobserving one feedback meeting\b",
    r"\bpresent on \d{1,2}\s+[A-Za-z]{3,9}\b",
    r"\battendance to coaching sessions run by\b",
    r"\bperformance checkpoint note for\b",
    r"\bfrom Mr\.\s*Brooks on two occasions\b.*?while looking",
    r"\baimed at the group,?\s*while looking\b",
    r"\bbrusque/ill-judged public remarks\b",
    r"\braised voice during feedback\b",
    r"\bsarcastic written tone and non-verbal disrespect\b",
    r"\bpublic challenge of instructions\b",
    r"\bthrough 1:1 or post-brief\b",
    r"\bprofessional influence\b",
    r"\bleading difficult conversations\b",
    r"\bWitness \d+ reports seeing eye-?rolling once\b",
    r"\bchairing remained with Mr\.\s*\w+\b",
    r"Handover log notes Ms\.\s*\w+\s+\"remained on Lane \d.*?coaching\.?\"",
    r"Ms\.\s*\w+\s+states she proposed a 15-minute delay.*?screen\.?",
    r"Witness \d+ recalls one instance of Ms\.\s*\w+.*?handover\.?",
    r"Mr\.\s*Brooks'?\s*communication style in some briefings fell below",
    r"Mr\.\s*Brooks to adopt SBII feedback.*?single out individuals\.?",
    r"Ms\.\s*Ahmed:?\s*sarcastic written tone.*?instructions\.?",
    r"Ms\.\s*Ahmed to channel challenges.*?safety-critical\.?",
    r"Letter to Mr\.\s*Brooks:.*?right to appeal\.?",
    r"Letter to Ms\.\s*Ahmed:.*?right to appeal\.?",
    r"Leading Difficult Conversations.*?6 weeks\.?",
    r"Professional Influence & Upward Challenge.*?6 weeks\.?",
    r"\d{1,2}\s+Jan\s+(?:AM|PM):\s*(?:Ms\.|Mr\.)\s*\w+\s*\(with\s+\w+.*?\)",
    r"\bManager 1:1 with Sophie by \d{1,2} Jan\b",
    r"\bCapture strengths and \d[–-]\d development commitments\b.*?\)\.?",
    r"\bAlign recognition:\s*Spot Award nomination recommended\.?",
    r"\bPayroll/Comp will align with calibration.*?conversation\)\.?",
    r"\bFile all evidence and the panel rationale.*?folder\.?",
    r"\busually cross-checkpoint,?\s*sustained quarters.*?delta\)\??",
    r"\bOccupational Health\b.*?\breasonable adjustments\b",
    r"\bself-referral regarding prolonged standing\b",
    r"\bfit for full duties\b",
    r"\bNo medical diagnosis disclosed\b",
    r"\bcompassionate leave\b.*?\bbereavement\b",
    r"\bI am submitting a formal complaint\b",
    r"\bmy colleague\b",
    r"\bwhere I felt undermined in front of the team\b",
    r"\bHappy to provide dates,?\s*witnesses,?\s*and Teams messages\.?",
    r"\bI felt singled out\b",
    r"\bmy area;?\s*I felt singled out\b",
    r"\bmy pace\b.*?\bholding the lane\b",
    r"\bI had been supporting a PRM assist\b",
    r"\bI believe\b.*?\bsarcastic\b",
    r"\bI'?m comfortable with the informal plan\b",
    r"\bAgree briefing language that avoids remarks\b.*?\btargeted\b",
    r"\bSet tone expectations at repack during peaks\b",
    r"\bImprove transparency on lane assignment\b.*?\bspikes\b",
    r"\bDesired outcomes:\b",
    r"\bThanks for raising this\b.*?\bneed-to-know basis\b",
    r"\b1\)\s*Initial fact-finding\b.*?\bconversations\b",
    r"\b2\)\s*Neutral evidence review\b.*?\bcontext\b",
    r"\b3\)\s*Consider informal resolution\b.*?\bneeded\b",
    r"\bPlease send your timeline\b.*?\bguidance\b",
    r"\bThanks for coming forward\b.*?\bmeeting\b",
    r"\bDates/times/locations\b.*?\bimpact\b",
    r"\bAny witnesses and Teams/email references\b",
    r"\bWhat outcome you'?re seeking\b.*?\btraining\b",
    r"\bWe'?ll apply Dignity at Work principles\b.*?\bmeeting\b",
    r"\bShared to your secure folder\b",
    r"\bBriefing bullets for the dates\b.*?\brecorded\b",
    r"\bRota extracts \+ note\b.*?\b12:30\b",
    r"\bReminder:\s*keep distribution tight\b.*?\bdetails\b",
    r"\bOK to share minimal\b.*?\bparties \+ HR\b",
    r"\bYes,?\s*issue a simple preservation notice\b.*?\badvised\b",
    r"\bWe'?ll support if it proceeds\b.*?\bDSAR\b",
    r"\bJenna:\s*Felt singled out\b.*?\bsarcastic\b",
    r"\bBriefing bullets don'?t capture phrasing\b",
    r"\bAddress via informal resolution\b.*?\bmediation\b",
    r"\bConduct reminders to both\b.*?\bsingle out\b",
    r"\bOptional mediation\b.*?\bworking\b",
    r"\bShort action plan\b.*?\bweeks\b",
    r"\bKeep preservation in place\b.*?\bfile\b",
    r"\bPlease retain materials relevant\b.*?\bhold\b",
    r"\bJenna:\s*\d{2}:\d{2}[–-]\d{2}:\d{2}\s*tomorrow\b.*?\bRoom\s+\w\b",
    r"\bPlease confirm availability for a 60-minute\b.*?\bmediation\b",
    r"\bOnce mediation concludes\b.*?\blifted\b",
]

DEFAULT_PRIVILEGED_KEYWORDS = [
    "legal review request",
    "legal confirms",
    "whether redaction",
    "whether any exemption applies",
    "preventing disclosure",
    "legal confirms the following",
    "third-party personal data must be redacted",
    "written summary may be provided instead",
    "legal counsel",
    "disclosability",
    "legal guidance request",
    "disclosure scope",
    "preservation",
    "redacting third-party names",
    "redacting third-party personal data",
    "strictly necessary snippets",
    "distribution limited to parties",
    "keep distribution tight",
    "remove any non-essential third-party details",
    "we'll support if it proceeds to formal er or dsar",
]

CASE_KEYWORDS = [
    "dsar", "data subject access", "subject access", "sar",
    "investigation", "incident", "claim", "complaint", "case",
    "cctv", "footage", "vehicle", "registration", "parking",
    "damage", "reference", "id", "photo id", "v5c",
    "grievance", "disciplinary", "bullying", "harassment",
    "witness", "interview", "evidence", "appeal", "outcome",
    "dignity at work", "fact-finding", "mediation",
    "conduct reminder", "preservation notice",
    "informal resolution", "formal er",
]

SECTION_HINTS = {
    "findings", "analysis", "recommendations", "appendix", "appendices",
    "evidence", "evidence considered", "interview", "chronology",
    "action plan", "recommended outcome letters", "communications plan",
    "risk assessment", "risk assessment & mitigations", "process summary",
    "parties", "terms of reference", "allegations summary",
    "witness statements", "operational records", "internal communications",
    "customer relations records", "cctv index", "processing information",
    "supervisor log", "handover notes", "full incident investigation report",
    "timeline of events", "executive summary",
    "data protection & disclosure notes", "sign-off",
    "attendance & sign-off", "allegation-to-evidence matrix",
    "interview schedule", "calibration", "improvement plan",
    "monitoring & support", "management actions", "training",
}

SENSITIVE_ROLE_WORDS = [
    "investigating manager", "appointing manager", "hr advisor",
    "hr adviser", "hr business partner", "hrbp", "hr coordinator",
    "hr coord", "hr shared services", "security duty manager",
    "operations manager", "ops manager", "team leader", "senior screener",
    "checkpoint lead", "security shift supervisor", "shift supervisor",
    "shift lead", "duty manager", "employee companion", "union rep",
    "union representative", "trade union representative", "companion",
    "new starter", "witness", "chair", "note-taker", "note taker",
    "investigating manager (im)", "queue marshal",
    "finance business partner", "finance bp", "panel member",
    "security compliance manager", "customer experience lead",
    "calibration chair", "legal counsel",
    "officer #1", "officer #2", "officer #3", "officer #4",
    "checkpoint supervisor", "ops mgr",
    "hr (emma w.)", "ops mgr (priya s.)",
]

TRAINING_OR_ACTION_HINTS = [
    "leading difficult conversations", "professional influence",
    "professional influence & upward challenge", "upward challenge",
    "facilitated mediation", "communication standards reset",
    "written conduct reminder", "teams/email etiquette refresher",
    "shadowing/coaching protocol", "coaching/rotation protocol",
    "action plan", "owner", "due", "letter to",
    "policy refresher", "contact protocol",
    "attendance improvement plan", "trigger awareness",
    "fortnightly check-ins", "coaching for officers",
    "reinforcement briefing", "spot observations", "courteous phrasing",
    "manager 1:1 with", "capture strengths", "development commitments",
    "spot award nomination", "formalise playbook",
    "cross-checkpoint training", "data storytelling",
    "sbii feedback technique", "situation-behaviour-impact-intent",
    "de-escalation refresher", "huddle language standard",
    "pulse feedback", "joint check-ins", "mediation session",
    "conduct reminders", "neutral briefing scripts",
    "preservation notice", "close-out emails",
]


# ============================================================
# REGEX / GLOBALS
# ============================================================
EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
PHONE_RE = re.compile(r"(?:(?:\+?\d{1,3}\s*)?(?:\(?0?\d{2,4}\)?\s*)?\d{3,4}\s*\d{3,4})")
HEADER_RE = re.compile(r"(?i)^\s*(from|to|cc|bcc|subject|date)\s*:")
HEADER_FIELD_RE = re.compile(r"(?i)^\s*(from|to|cc|bcc|subject|date)\s*:\s*(.*)$")

SIGNOFF_RE = re.compile(
    r"(?i)^\s*(kind regards|regards|best regards|best|many thanks|thanks|thank you|"
    r"yours sincerely|yours faithfully|sincerely|cheers)\s*[,!:]?\s*$"
)
MOBILE_FOOTER_RE = re.compile(r"(?i)^\s*sent from my .+\s*$")

ROLE_LINE_RE = re.compile(
    r"^(Employee|Chair|HR Advisor|HR Adviser|Investigating Manager \(IM\)|Investigating Manager|"
    r"Note-Taker|Employee Companion|Appointing Manager|Witness \d+|Complainant A / Respondent B|"
    r"Complainant B / Respondent A)\s*:\s*(.+)$",
    re.IGNORECASE,
)
TITLE_RE = re.compile(r"^(Mr|Mrs|Ms|Miss|Dr)\.?\s+", re.IGNORECASE)
GREETING_RE = re.compile(r"(?i)^\s*(dear|hi)\s+([^,]+),\s*$")
DATE_WORD_RE = re.compile(r"\b\d{1,2}\s+[A-Za-z]{3,9}\b")
WITNESS_LINE_RE = re.compile(r"(?i)^\s*(witness\s*\d+\s*:\s*)(.+)$")
INTERVIEW_LINE_RE = re.compile(
    r"(?i)^\s*(?:[-•]\s*)?(?:\d{1,2}\s+[A-Za-z]{3,9}(?:\s*(?:AM|PM))?\s*:\s*)(.+)$"
)
LETTER_LINE_RE = re.compile(r"(?i)(letter to\s+)([^:]+)(:)")
OWNER_LINE_RE = re.compile(r"(?i)\b(owner)\b")
ACTION_OWNER_ROW_HINT_RE = re.compile(r"(?i)\b(action|owner|due)\b")
FULL_PERSON_RE = re.compile(r"\b(?:Mr|Mrs|Ms|Miss|Dr)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?\b")
SIMPLE_NAME_RE = re.compile(r"\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\b")
BRACKET_CONTENT_RE = re.compile(r"\(([^)]+)\)")
INLINE_AGAINST_RE = re.compile(r"(?i)\bagainst\s+([^.,;]+)")
ROLE_NAME_COMBO_RE = re.compile(
    r"(?i)\b(?:investigating manager|appointing manager|hr advisor|hr adviser|ops mgr|owner)\s*[:\-]?\s*"
    r"([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,2})"
)


# ============================================================
# CLICKABLE PREVIEW COMPONENT
# ============================================================
COMPONENT_NAME = "clickable_redaction_preview"
COMPONENT_DIR = Path(__file__).resolve().parent / COMPONENT_NAME
COMPONENT_INDEX = COMPONENT_DIR / "index.html"


def clickable_preview_available() -> bool:
    return COMPONENT_INDEX.is_file()


def _jsonify_candidate(candidate: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "id": str(candidate.get("id", "")),
        "para_idx": int(candidate.get("para_idx", 0)),
        "start": int(candidate.get("start", 0)),
        "end": int(candidate.get("end", 0)),
        "text": str(candidate.get("text", "")),
        "detector": str(candidate.get("detector", "")),
        "entity_label": str(candidate.get("entity_label", "")),
        "reason": str(candidate.get("reason", "")),
        "suggested_decision": str(candidate.get("suggested_decision", "review")),
        "decision": str(candidate.get("decision", candidate.get("suggested_decision", "review"))),
    }


if clickable_preview_available():
    _CLICKABLE_PREVIEW = components.declare_component(
        COMPONENT_NAME,
        path=str(COMPONENT_DIR),
    )
else:
    _CLICKABLE_PREVIEW = None


def clickable_redaction_preview(
    paragraphs: List[str],
    candidates: List[Dict[str, Any]],
    key: str,
    default: Dict[str, Any] | None = None,
):
    if _CLICKABLE_PREVIEW is None:
        return default

    safe_paragraphs = [str(p) for p in (paragraphs or [])]
    safe_candidates = [_jsonify_candidate(c) for c in (candidates or [])]

    return _CLICKABLE_PREVIEW(
        paragraphs=safe_paragraphs,
        candidates=safe_candidates,
        key=key,
        default=default or {},
    )

# ============================================================
# MODEL LOADING
# ============================================================
@st.cache_resource(show_spinner=False)
def load_ner_pipeline():
    device = 0 if torch.cuda.is_available() else -1
    return pipeline(
        "ner",
        model="dslim/bert-base-NER",
        aggregation_strategy="simple",
        device=device,
    )


# ============================================================
# HELPERS
# ============================================================
def normalize(s: str) -> str:
    return re.sub(r"\s+", " ", s).strip()


def normalize_quotes(s: str) -> str:
    return s.replace("\u2018", "'").replace("\u2019", "'").replace("\u201c", '"').replace("\u201d", '"').replace("\u2013", "-").replace("\u2014", "-")


def paragraph_lines(p: str) -> List[str]:
    return p.splitlines() if "\n" in p else [p]


def paragraph_lines_with_offsets(p: str) -> List[Tuple[str, int, int]]:
    if "\n" not in p:
        return [(p, 0, len(p))]
    out = []
    pos = 0
    for chunk in p.splitlines(keepends=True):
        line = chunk.rstrip("\r\n")
        out.append((line, pos, pos + len(line)))
        pos += len(chunk)
    return out


def starts_new_email(paragraph_text: str) -> bool:
    for ln in paragraph_lines(paragraph_text):
        if re.match(r"(?i)^\s*from\s*:", ln.strip()):
            return True
    return False


def split_into_blocks(paras: List[str]) -> List[List[str]]:
    blocks: List[List[str]] = []
    cur: List[str] = []
    for p in paras:
        if starts_new_email(p) and cur:
            blocks.append(cur)
            cur = [p]
        else:
            cur.append(p)
    if cur:
        blocks.append(cur)
    return blocks


def merge_spans(spans: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
    if not spans:
        return []
    spans = sorted(spans)
    merged = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    return [(a, b) for a, b in merged]


def redact_by_spans(text: str, spans: List[Tuple[int, int]], token: str = REDACT_TOKEN) -> str:
    if not spans:
        return text
    merged = merge_spans(spans)
    out = []
    prev = 0
    for s, e in merged:
        out.append(text[prev:s])
        out.append(token)
        prev = e
    out.append(text[prev:])
    return "".join(out)


def find_email_spans(text: str) -> List[Tuple[int, int, str]]:
    return [(m.start(), m.end(), m.group(0)) for m in EMAIL_RE.finditer(text)]


def find_phone_spans(text: str) -> List[Tuple[int, int, str]]:
    out = []
    for m in PHONE_RE.finditer(text):
        s = m.group(0).strip()
        digits = re.sub(r"\D", "", s)
        if len(digits) >= 9:
            out.append((m.start(), m.end(), m.group(0)))
    return out


def find_person_spans(text: str, ner_pipe) -> List[Tuple[int, int, str, str]]:
    ents = ner_pipe(text) or []
    out = []
    for e in ents:
        if e.get("entity_group") == "PER":
            start, end = int(e["start"]), int(e["end"])
            out.append((start, end, text[start:end], "PER"))
    return out


def add_candidate(
    candidates: List[Dict[str, Any]],
    seen: set,
    *,
    para_idx: int,
    start: int,
    end: int,
    text: str,
    detector: str,
    entity_label: str,
    reason: str,
    suggested_decision: str = "redact",
):
    if start < 0 or end <= start:
        return
    if end > len(text):
        end = len(text)
    key = (para_idx, start, end, detector, entity_label)
    if key in seen:
        return
    seen.add(key)
    candidates.append(
        {
            "id": f"C{len(candidates)+1:05d}",
            "para_idx": para_idx,
            "start": start,
            "end": end,
            "text": text[start:end],
            "detector": detector,
            "entity_label": entity_label,
            "reason": reason,
            "suggested_decision": suggested_decision,
            "decision": suggested_decision,
        }
    )


def literal_occurrences(text: str, needle: str) -> List[Tuple[int, int]]:
    out = []
    if not needle:
        return out
    start = 0
    while True:
        i = text.find(needle, start)
        if i == -1:
            break
        out.append((i, i + len(needle)))
        start = i + len(needle)
    return out


def overlap(a: Tuple[int, int], b: Tuple[int, int]) -> bool:
    return not (a[1] <= b[0] or a[0] >= b[1])


def build_subject_patterns(subject_full: str, subject_email: str, aliases: List[str]) -> List[re.Pattern]:
    pats: List[re.Pattern] = []
    values = [subject_full, subject_email] + list(aliases)
    for v in values:
        v = v.strip()
        if v:
            pats.append(re.compile(rf"(?i)\b{re.escape(v)}\b"))
    if subject_full.strip():
        parts = subject_full.strip().split()
        last = parts[-1]
        pats.append(re.compile(rf"(?i)\b(mr|mrs|ms|miss|dr)\.?\s+{re.escape(last)}\b"))
        if len(parts) >= 2:
            pats.append(re.compile(rf"(?i)\b{re.escape(parts[0])}\b"))
            pats.append(re.compile(rf"(?i)\b{re.escape(last)}\b"))
    if subject_email.strip() and "@" in subject_email:
        local = subject_email.split("@", 1)[0]
        pats.append(re.compile(rf"(?i)\b{re.escape(local)}\b"))
    return pats


def mentions_subject(text: str, subject_pats: List[re.Pattern]) -> bool:
    t = text.strip()
    if not t:
        return False
    return any(p.search(t) for p in subject_pats)


def is_allowlisted_email(email: str, keep_emails_lower: set, subject_email_lower: str) -> bool:
    el = email.lower()
    return el == subject_email_lower or el in keep_emails_lower


def is_case_related(text: str) -> bool:
    low = text.lower()
    return any(k in low for k in CASE_KEYWORDS)


def is_internal_email(email: str, domain: str = "heathrow.com") -> bool:
    return email.lower().endswith(f"@{domain}")


def localpart(email: str) -> str:
    return email.split("@", 1)[0]


def compile_patterns(lines: List[str]) -> List[re.Pattern]:
    out = []
    for line in lines:
        s = line.strip()
        if not s:
            continue
        try:
            out.append(re.compile(s, re.IGNORECASE))
        except re.error:
            out.append(re.compile(re.escape(s), re.IGNORECASE))
    return out


def signature_start_index(block_paras: List[str], sender_is_subject: bool) -> int:
    if sender_is_subject:
        return len(block_paras)
    scan_limit = min(len(block_paras), 12)
    start = len(block_paras)
    for i in range(len(block_paras) - 1, len(block_paras) - scan_limit - 1, -1):
        if i < 0:
            break
        t = block_paras[i].strip()
        if not t:
            continue
        if SIGNOFF_RE.match(t) or MOBILE_FOOTER_RE.match(t):
            start = i
            break
    if start == len(block_paras) and len(block_paras) >= 2:
        last = block_paras[-1].strip()
        prev = block_paras[-2].strip()
        if SIGNOFF_RE.match(prev) and 1 <= len(last.split()) <= 8:
            start = len(block_paras) - 2
    return start


def dedupe_overlaps_prefer_rules(candidates: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    by_para: Dict[int, List[Dict[str, Any]]] = {}
    for c in candidates:
        by_para.setdefault(c["para_idx"], []).append(c)
    priority = {"MANUAL": -1, "RULE": 0, "REGEX": 1, "NER": 2}
    out = []
    for para_idx in sorted(by_para):
        para_cands = sorted(
            by_para[para_idx],
            key=lambda c: (c["start"], priority.get(c["detector"], 9), -(c["end"] - c["start"])),
        )
        accepted = []
        occupied = []
        for c in para_cands:
            s, e = c["start"], c["end"]
            overlaps = any(not (e <= os or s >= oe) for os, oe in occupied)
            if overlaps:
                continue
            accepted.append(c)
            occupied.append((s, e))
        out.extend(sorted(accepted, key=lambda c: c["start"]))
    return out


def get_last_nonempty_heading(paras: List[str], para_idx: int) -> str:
    for i in range(para_idx, -1, -1):
        t = normalize(paras[i])
        if not t:
            continue
        low = t.lower()
        if low.endswith(":") or re.match(r"^\d+(\.\d+)?\)", low) or any(h in low for h in SECTION_HINTS):
            return t
    return ""


def line_has_sensitive_role_words(text: str) -> bool:
    low = text.lower()
    return any(w in low for w in SENSITIVE_ROLE_WORDS)


def line_has_training_or_action_hint(text: str) -> bool:
    low = text.lower()
    return any(w in low for w in TRAINING_OR_ACTION_HINTS)


def add_literal_if_non_subject(*, para_text, para_idx, needle, subject_pats, candidates, seen, entity_label, reason, suggested_decision="redact"):
    if not needle or mentions_subject(needle, subject_pats):
        return
    for s, e in literal_occurrences(para_text, needle):
        add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=para_text, detector="RULE", entity_label=entity_label, reason=reason, suggested_decision=suggested_decision)


def add_full_match_if_non_subject(*, para_text, para_idx, pattern, subject_pats, candidates, seen, entity_label, reason, group_index=0, suggested_decision="redact"):
    for m in pattern.finditer(para_text):
        target = m.group(group_index)
        if not target or mentions_subject(target, subject_pats):
            continue
        add_candidate(candidates, seen, para_idx=para_idx, start=m.start(group_index), end=m.end(group_index), text=para_text, detector="RULE", entity_label=entity_label, reason=reason, suggested_decision=suggested_decision)


# ============================================================
# TRANSCRIPT POLICY
# ============================================================
def parse_name(full_name: str) -> Dict[str, str]:
    raw = normalize_quotes(full_name.strip())
    title = ""
    nickname = ""
    m = TITLE_RE.match(raw)
    if m:
        title = m.group(1)
        raw = raw[m.end():].strip()
    nick_m = re.search(r'"([^"]+)"', raw)
    if nick_m:
        nickname = nick_m.group(1).strip()
        raw_wo_nick = re.sub(r'"[^"]+"', "", raw).strip()
    else:
        raw_wo_nick = raw
    parts = raw_wo_nick.split()
    first = parts[0] if parts else ""
    last = parts[-1] if len(parts) >= 2 else (parts[0] if parts else "")
    return {"title": title, "first": first, "nickname": nickname, "last": last, "full_name": full_name.strip()}


def extract_transcript_policy(paras: List[str]) -> Dict[str, Any]:
    subject = None
    participants: List[Dict[str, Any]] = []
    for p in paras:
        m = ROLE_LINE_RE.match(p.strip())
        if not m:
            continue
        role = m.group(1).strip()
        value = m.group(2).strip()
        if role.lower() == "employee":
            name_part = value.split("(Employee ID", 1)[0].strip()
            subject = {"role": role, "raw_value": value, "name_part": name_part, **parse_name(name_part)}
        else:
            name_part = value.split(",", 1)[0].strip()
            participants.append({"role": role, "raw_value": value, "name_part": name_part, **parse_name(name_part)})
    if subject is None:
        subject = {"role": "Employee", "raw_value": "", "name_part": "", "title": "", "first": "", "nickname": "", "last": "", "full_name": ""}
    subject_aliases: Set[str] = set()
    for x in [subject["full_name"], subject["first"], subject["nickname"]]:
        if x and x.strip():
            subject_aliases.add(x.strip())
    if subject["last"]:
        subject_aliases.add(subject["last"])
        if subject["title"]:
            subject_aliases.add(f"{subject['title']} {subject['last']}")
            subject_aliases.add(f"{subject['title']}. {subject['last']}")
    third_party_full_names = sorted({p["full_name"] for p in participants if p["full_name"]}, key=len, reverse=True)
    third_party_surnames = sorted({p["last"] for p in participants if p["last"]}, key=len, reverse=True)
    third_party_header_values = [(p["role"], p["raw_value"]) for p in participants if p["raw_value"]]
    return {"subject": subject, "subject_aliases": sorted(subject_aliases, key=len, reverse=True), "participants": participants, "third_party_full_names": third_party_full_names, "third_party_surnames": third_party_surnames, "third_party_header_values": third_party_header_values}


def collect_transcript_policy_candidates(*, paras, candidates, seen, subject_pats, transcript_policy, private_patterns):
    surname_set = {normalize_quotes(x).lower() for x in transcript_policy.get("third_party_surnames", [])}
    for para_idx, p in enumerate(paras):
        if not p.strip():
            continue
        m = ROLE_LINE_RE.match(p.strip())
        if m:
            role = m.group(1).strip()
            if role.lower() != "employee":
                prefix = p[: p.lower().find(":") + 1]
                line_start = len(prefix)
                while line_start < len(p) and p[line_start].isspace():
                    line_start += 1
                add_candidate(candidates, seen, para_idx=para_idx, start=line_start, end=len(p), text=p, detector="RULE", entity_label="HEADER_PARTICIPANT", reason=f"Transcript header non-subject participant: {role}")
        for full_name in transcript_policy.get("third_party_full_names", []):
            for s, e in literal_occurrences(p, full_name):
                if mentions_subject(p[s:e], subject_pats):
                    continue
                add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=p, detector="RULE", entity_label="THIRD_PARTY_NAME", reason="Third-party participant full name")
        for m in re.finditer(r"\[([^\]]+)\]", p):
            tag = normalize_quotes(m.group(1).strip()).lower()
            if tag in surname_set:
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(1), end=m.end(1), text=p, detector="RULE", entity_label="SPEAKER_TAG_SURNAME", reason="Third-party speaker tag surname in brackets")
        for m in re.finditer(r"\(([^)]+)\)\s*:", p):
            tag = normalize_quotes(m.group(1).strip()).lower()
            if tag in surname_set:
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(1), end=m.end(1), text=p, detector="RULE", entity_label="SPEAKER_TAG_SURNAME", reason="Third-party speaker tag surname in parentheses")
        for pat in private_patterns:
            for m in pat.finditer(p):
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(), end=m.end(), text=p, detector="RULE", entity_label="PRIVATE_SNIPPET", reason=f"Matched transcript private pattern: {pat.pattern}")


# ============================================================
# EMAIL POLICY
# ============================================================
def parse_email_header_line(line_text: str) -> Tuple[str, str] | None:
    m = HEADER_FIELD_RE.match(line_text)
    if not m:
        return None
    return m.group(1).lower(), m.group(2)


def derive_name_tokens_from_email(email: str) -> Set[str]:
    tokens: Set[str] = set()
    lp = localpart(email).replace("_", ".").replace("-", ".")
    parts = [x for x in lp.split(".") if x]
    if not parts:
        return tokens
    if len(parts) >= 2:
        first = parts[0]
        last = parts[-1]
        tokens.add(first.lower())
        tokens.add(last.lower())
        tokens.add(f"{first.lower()} {last.lower()}")
    else:
        tokens.add(parts[0].lower())
    return tokens


def extract_email_block_context(block_paras, subject_pats, subject_email_lower, keep_emails_lower, privileged_keywords):
    header_emails = []
    sender_emails = []
    recipient_emails = []
    staff_name_tokens: Set[str] = set()
    all_lines = []
    for p in block_paras:
        for ln in paragraph_lines(p):
            line = ln.strip()
            if not line:
                continue
            all_lines.append(line)
            parsed = parse_email_header_line(line)
            if not parsed:
                continue
            field, value = parsed
            emails = [m.group(0).lower() for m in EMAIL_RE.finditer(value)]
            header_emails.extend(emails)
            if field == "from":
                sender_emails.extend(emails)
            elif field in {"to", "cc", "bcc"}:
                recipient_emails.extend(emails)
            for em in emails:
                if em == subject_email_lower:
                    continue
                if is_allowlisted_email(em, keep_emails_lower, subject_email_lower):
                    continue
                if is_internal_email(em):
                    staff_name_tokens |= derive_name_tokens_from_email(em)
    sender_is_subject = any(em == subject_email_lower for em in sender_emails) or any(mentions_subject(line, subject_pats) for line in all_lines if line.lower().startswith("from:"))
    addressed_to_subject = any(em == subject_email_lower for em in recipient_emails) or any(mentions_subject(line, subject_pats) for line in all_lines if re.match(r"(?i)^\s*(to|cc|bcc)\s*:", line))
    subject_in_headers = sender_is_subject or addressed_to_subject
    subject_anywhere = any(mentions_subject(line, subject_pats) for line in all_lines)
    internal_only = len(header_emails) > 0 and all(is_internal_email(em) for em in header_emails) and not subject_in_headers
    block_text_low = "\n".join(block_paras).lower()
    legal_internal = internal_only and any(k.lower() in block_text_low for k in privileged_keywords)
    return {"sender_is_subject": sender_is_subject, "addressed_to_subject": addressed_to_subject, "subject_in_headers": subject_in_headers, "subject_anywhere": subject_anywhere, "internal_only": internal_only, "legal_internal": legal_internal, "staff_name_tokens": sorted(staff_name_tokens)}


def collect_email_header_candidates(*, para_text, para_idx, line_text, line_start, subject_email_lower, keep_emails_lower, subject_in_headers, candidates, seen):
    parsed = parse_email_header_line(line_text)
    if not parsed:
        return
    field, value = parsed
    if field not in {"from", "to", "cc", "bcc"}:
        return
    colon_pos = line_text.find(":")
    if colon_pos == -1:
        return
    for m in EMAIL_RE.finditer(line_text[colon_pos + 1:]):
        email = m.group(0)
        email_low = email.lower()
        abs_s = line_start + colon_pos + 1 + m.start()
        abs_e = line_start + colon_pos + 1 + m.end()
        if email_low == subject_email_lower:
            continue
        if subject_in_headers and is_allowlisted_email(email_low, keep_emails_lower, subject_email_lower):
            continue
        if is_internal_email(email_low):
            lp = localpart(email)
            add_candidate(candidates, seen, para_idx=para_idx, start=abs_s, end=abs_s + len(lp), text=para_text, detector="RULE", entity_label="INTERNAL_EMAIL_LOCALPART", reason=f"Internal non-subject email localpart in {field.upper()} header")
        else:
            add_candidate(candidates, seen, para_idx=para_idx, start=abs_s, end=abs_e, text=para_text, detector="RULE", entity_label="EMAIL", reason=f"Non-subject external email in {field.upper()} header")


def collect_email_greeting_candidates(*, para_text, para_idx, line_text, line_start, subject_pats, staff_name_tokens, candidates, seen):
    m = GREETING_RE.match(line_text.strip())
    if not m:
        return
    addressee = normalize_quotes(m.group(2).strip())
    if mentions_subject(addressee, subject_pats):
        return
    add_low = addressee.lower()
    generic_team = any(word in add_low for word in ["team", "support", "security", "legal", "dsar"])
    if generic_team and add_low not in {t.lower() for t in staff_name_tokens}:
        return
    if staff_name_tokens and add_low not in {t.lower() for t in staff_name_tokens}:
        if not re.search(r"[A-Za-z]", addressee):
            return
    rel_s = line_text.lower().find(addressee.lower())
    if rel_s == -1:
        return
    rel_e = rel_s + len(addressee)
    add_candidate(candidates, seen, para_idx=para_idx, start=line_start + rel_s, end=line_start + rel_e, text=para_text, detector="RULE", entity_label="EMAIL_GREETING_NAME", reason="Greeting to non-subject person")


def collect_regex_candidates_in_para(*, para_text, para_idx, subject_email_lower, keep_emails_lower, candidates, seen, reason_prefix):
    for s, e, email in find_email_spans(para_text):
        if is_allowlisted_email(email, keep_emails_lower, subject_email_lower):
            continue
        add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=para_text, detector="REGEX", entity_label="EMAIL", reason=f"{reason_prefix}; email")
    for s, e, _ in find_phone_spans(para_text):
        add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=para_text, detector="REGEX", entity_label="PHONE", reason=f"{reason_prefix}; phone")


def collect_email_chain_candidates(*, paras, subject_pats, subject_email_lower, keep_emails_lower, privileged_keywords, private_patterns, enable_regex, enable_ner_backstop, ner_pipe, candidates, seen):
    blocks = split_into_blocks(paras)
    para_global_idx = 0
    block_meta = []
    for block_idx, block in enumerate(blocks):
        ctx = extract_email_block_context(block, subject_pats=subject_pats, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, privileged_keywords=privileged_keywords)
        sig_start = signature_start_index(block, sender_is_subject=ctx["sender_is_subject"])
        block_meta.append({"block_idx": block_idx, "sender_is_subject": ctx["sender_is_subject"], "addressed_to_subject": ctx["addressed_to_subject"], "subject_in_headers": ctx["subject_in_headers"], "subject_anywhere": ctx["subject_anywhere"], "internal_only": ctx["internal_only"], "legal_internal": ctx["legal_internal"], "sig_start": sig_start, "staff_name_tokens": ctx["staff_name_tokens"], "para_start_idx": para_global_idx, "para_count": len(block)})
        if ctx["legal_internal"]:
            for i, p in enumerate(block):
                current_para_idx = para_global_idx + i
                if not p.strip():
                    continue
                add_candidate(candidates, seen, para_idx=current_para_idx, start=0, end=len(p), text=p, detector="RULE", entity_label="INTERNAL_PRIVILEGED_BLOCK", reason="Internal privileged/legal review block")
            para_global_idx += len(block)
            continue
        for i, p in enumerate(block):
            current_para_idx = para_global_idx + i
            if not p.strip():
                continue
            if (i >= sig_start) and (not ctx["sender_is_subject"]):
                add_candidate(candidates, seen, para_idx=current_para_idx, start=0, end=len(p), text=p, detector="RULE", entity_label="EMAIL_SIGNATURE_BLOCK", reason="Non-subject email signature / sign-off block")
                continue
            for line_text, line_start, line_end in paragraph_lines_with_offsets(p):
                s = line_text.strip()
                if not s:
                    continue
                parsed = parse_email_header_line(s)
                if parsed:
                    field = parsed[0]
                    if field in {"from", "to", "cc", "bcc"}:
                        collect_email_header_candidates(para_text=p, para_idx=current_para_idx, line_text=line_text, line_start=line_start, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, subject_in_headers=ctx["subject_in_headers"], candidates=candidates, seen=seen)
                        continue
                collect_email_greeting_candidates(para_text=p, para_idx=current_para_idx, line_text=line_text, line_start=line_start, subject_pats=subject_pats, staff_name_tokens=ctx["staff_name_tokens"], candidates=candidates, seen=seen)
                for pat in private_patterns:
                    for m in pat.finditer(line_text):
                        add_candidate(candidates, seen, para_idx=current_para_idx, start=line_start + m.start(), end=line_start + m.end(), text=p, detector="RULE", entity_label="PRIVATE_SNIPPET", reason=f"Matched private pattern: {pat.pattern}")
            if enable_regex:
                collect_regex_candidates_in_para(para_text=p, para_idx=current_para_idx, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, candidates=candidates, seen=seen, reason_prefix="Email policy")
            if enable_ner_backstop and ner_pipe is not None:
                try:
                    persons = find_person_spans(p, ner_pipe)
                except Exception:
                    persons = []
                for s, e, ent_text, label in persons:
                    if mentions_subject(ent_text, subject_pats):
                        continue
                    add_candidate(candidates, seen, para_idx=current_para_idx, start=s, end=e, text=p, detector="NER", entity_label=label, reason="NER backstop for non-subject person in email chain", suggested_decision="review")
        para_global_idx += len(block)
    return block_meta


# ============================================================
# CASE DOCUMENT POLICY
# ============================================================
def collect_case_document_policy_candidates(*, paras, subject_pats, transcript_policy, private_patterns, candidates, seen):
    third_party_full_names = transcript_policy.get("third_party_full_names", [])
    third_party_surnames = transcript_policy.get("third_party_surnames", [])
    participants = transcript_policy.get("participants", [])
    known_non_subject_strings: Set[str] = set()
    for p in participants:
        for v in [p.get("full_name", ""), p.get("first", ""), p.get("last", "")]:
            if v and not mentions_subject(v, subject_pats):
                known_non_subject_strings.add(v)
    for para_idx, p in enumerate(paras):
        t = normalize_quotes(p)
        low = t.lower()
        heading = get_last_nonempty_heading(paras, para_idx).lower()
        if not t:
            continue
        if ":" in t and line_has_sensitive_role_words(t):
            value = t.split(":", 1)[1].strip()
            if value and not mentions_subject(value, subject_pats):
                add_candidate(candidates, seen, para_idx=para_idx, start=t.find(value), end=t.find(value) + len(value), text=t, detector="RULE", entity_label="ROLE_VALUE_NON_SUBJECT", reason="Role/value header identifies non-subject person")
        if re.search(r"(?i)(complainant|respondent)", t):
            if ":" in t:
                value = t.split(":", 1)[1].strip()
                if value and not mentions_subject(value, subject_pats):
                    add_candidate(candidates, seen, para_idx=para_idx, start=t.find(value), end=t.find(value) + len(value), text=t, detector="RULE", entity_label="PARTY_LINE_NON_SUBJECT", reason="Party line identifies non-subject party")
        wm = WITNESS_LINE_RE.match(t)
        if wm:
            value = wm.group(2).strip()
            if value and not mentions_subject(value, subject_pats):
                add_candidate(candidates, seen, para_idx=para_idx, start=wm.start(2), end=wm.end(2), text=t, detector="RULE", entity_label="WITNESS_DESCRIPTOR", reason="Witness descriptor/date should be redacted")
        im = INTERVIEW_LINE_RE.match(t)
        if im:
            value = im.group(1).strip()
            if value and not mentions_subject(value, subject_pats):
                if FULL_PERSON_RE.search(value) or line_has_sensitive_role_words(value) or "with " in value.lower():
                    add_candidate(candidates, seen, para_idx=para_idx, start=im.start(1), end=im.end(1), text=t, detector="RULE", entity_label="INTERVIEW_SCHEDULE_ENTRY", reason="Interview schedule non-subject entry")
        for m in LETTER_LINE_RE.finditer(t):
            letter_target = m.group(2).strip()
            if letter_target and not mentions_subject(letter_target, subject_pats):
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(2), end=m.end(2), text=t, detector="RULE", entity_label="OUTCOME_LETTER_TARGET", reason="Outcome letter target is non-subject")
                tail_start = m.start(1)
                add_candidate(candidates, seen, para_idx=para_idx, start=tail_start, end=len(t), text=t, detector="RULE", entity_label="OUTCOME_LETTER_LINE", reason="Non-subject outcome letter line")
        for m in INLINE_AGAINST_RE.finditer(t):
            frag = m.group(1).strip().rstrip(").,;:")
            if frag and not mentions_subject(frag, subject_pats):
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(1), end=m.start(1) + len(frag), text=t, detector="RULE", entity_label="AGAINST_NON_SUBJECT", reason="Heading/allegation references non-subject person")
        for m in ROLE_NAME_COMBO_RE.finditer(t):
            target = m.group(1).strip()
            if target and not mentions_subject(target, subject_pats):
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(1), end=m.end(1), text=t, detector="RULE", entity_label="ROLE_NAME_COMBO", reason="Role line contains non-subject person")
        for full_name in sorted(known_non_subject_strings | set(third_party_full_names), key=len, reverse=True):
            if not full_name:
                continue
            for s, e in literal_occurrences(t, full_name):
                frag = t[s:e]
                if mentions_subject(frag, subject_pats):
                    continue
                redact_whole_clause = any(key in low for key in ["recommended outcome letters", "letter to", "action plan", "owner", "facilitated mediation", "communication standards reset", "training", "professional influence", "leading difficult conversations"])
                if redact_whole_clause:
                    add_candidate(candidates, seen, para_idx=para_idx, start=0, end=len(t), text=t, detector="RULE", entity_label="NON_SUBJECT_ACTION_LINE", reason="Action/training/outcome line tied to non-subject")
                else:
                    add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=t, detector="RULE", entity_label="THIRD_PARTY_NAME", reason="Non-subject name in case document")
        for m in FULL_PERSON_RE.finditer(t):
            frag = m.group(0)
            if mentions_subject(frag, subject_pats):
                continue
            add_candidate(candidates, seen, para_idx=para_idx, start=m.start(), end=m.end(), text=t, detector="RULE", entity_label="TITLED_PERSON_NAME", reason="Titled non-subject person mention")
        for m in BRACKET_CONTENT_RE.finditer(t):
            inner = m.group(1).strip()
            if not inner or mentions_subject(inner, subject_pats):
                continue
            if line_has_sensitive_role_words(inner) or DATE_WORD_RE.search(inner) or "with " in inner.lower() or "observ" in inner.lower() or "present" in inner.lower() or "lane" in inner.lower():
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(1), end=m.end(1), text=t, detector="RULE", entity_label="BRACKET_CONTEXT", reason="Bracketed witness/representation detail")
        for pat in private_patterns:
            for m in pat.finditer(t):
                snippet = m.group(0)
                if mentions_subject(snippet, subject_pats):
                    continue
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(), end=m.end(), text=t, detector="RULE", entity_label="PRIVATE_SNIPPET", reason=f"Matched case private pattern: {pat.pattern}")
        if line_has_training_or_action_hint(t):
            if any(not mentions_subject(name, subject_pats) and name in t for name in third_party_full_names):
                add_candidate(candidates, seen, para_idx=para_idx, start=0, end=len(t), text=t, detector="RULE", entity_label="ACTION_OR_TRAINING_LINE", reason="Training/action line tied to non-subject person")
        if ACTION_OWNER_ROW_HINT_RE.search(t) and line_has_sensitive_role_words(t):
            if any(not mentions_subject(name, subject_pats) and name in t for name in third_party_full_names):
                add_candidate(candidates, seen, para_idx=para_idx, start=0, end=len(t), text=t, detector="RULE", entity_label="ACTION_PLAN_ROW", reason="Action-plan row identifies non-subject")
        if "witness" in low and (DATE_WORD_RE.search(t) or "present" in low or "observ" in low or "lane" in low):
            if ":" in t:
                value = t.split(":", 1)[1].strip()
                if value and not mentions_subject(value, subject_pats):
                    add_candidate(candidates, seen, para_idx=para_idx, start=t.find(value), end=t.find(value) + len(value), text=t, detector="RULE", entity_label="WITNESS_DETAIL", reason="Witness detail should be redacted")
        if any(k in heading for k in ["recommend", "finding", "analysis", "evidence", "appendix", "chronology"]):
            for m in FULL_PERSON_RE.finditer(t):
                frag = m.group(0)
                if mentions_subject(frag, subject_pats):
                    continue
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(), end=m.end(), text=t, detector="RULE", entity_label="NON_SUBJECT_CONTEXT_PERSON", reason="Non-subject person singled out in sensitive section")


# ============================================================
# GENERIC / FALLBACK POLICY
# ============================================================
def collect_generic_candidates(*, paras, subject_pats, subject_email_lower, keep_emails_lower, private_patterns, privileged_keywords, enable_regex, enable_ner_backstop, ner_pipe, transcript_policy, candidates, seen):
    collect_case_document_policy_candidates(paras=paras, subject_pats=subject_pats, transcript_policy=transcript_policy, private_patterns=private_patterns, candidates=candidates, seen=seen)
    for para_idx, p in enumerate(paras):
        if not p.strip():
            continue
        low = p.lower()
        if not mentions_subject(p, subject_pats) and any(k in low for k in privileged_keywords):
            add_candidate(candidates, seen, para_idx=para_idx, start=0, end=len(p), text=p, detector="RULE", entity_label="GENERIC_PRIVILEGED", reason="Generic privileged / unrelated internal content")
            continue
        for pat in private_patterns:
            for m in pat.finditer(p):
                add_candidate(candidates, seen, para_idx=para_idx, start=m.start(), end=m.end(), text=p, detector="RULE", entity_label="PRIVATE_SNIPPET", reason=f"Matched generic private pattern: {pat.pattern}")
        if enable_regex:
            collect_regex_candidates_in_para(para_text=p, para_idx=para_idx, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, candidates=candidates, seen=seen, reason_prefix="Generic policy")
        if enable_ner_backstop and ner_pipe is not None:
            try:
                persons = find_person_spans(p, ner_pipe)
            except Exception:
                persons = []
            for s, e, ent_text, label in persons:
                if mentions_subject(ent_text, subject_pats):
                    continue
                add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=p, detector="NER", entity_label=label, reason="NER backstop for non-subject person", suggested_decision="review")


# ============================================================
# ANALYSIS DISPATCH
# ============================================================
def detect_doc_mode(paras: List[str]) -> str:
    role_lines = 0
    from_lines = 0
    for p in paras:
        if ROLE_LINE_RE.match(p.strip()):
            role_lines += 1
        for ln in paragraph_lines(p):
            if re.match(r"(?i)^\s*from\s*:", ln.strip()):
                from_lines += 1
    if role_lines >= 3:
        return "transcript"
    if from_lines >= 2:
        return "email"
    return "generic"


def analyze_docx_paragraphs(paras, *, subject_full, subject_email, aliases, keep_emails, private_patterns_text, privileged_keywords_text, enable_regex, enable_ner_backstop, auto_extract_subject, document_mode, ner_pipe=None):
    detected_mode = detect_doc_mode(paras)
    effective_mode = detected_mode if document_mode == "auto" else document_mode
    keep_emails_lower = {e.lower().strip() for e in keep_emails if e.strip()}
    subject_email_lower = subject_email.lower().strip()
    transcript_policy = extract_transcript_policy(paras)
    extracted_subject_full = transcript_policy["subject"].get("full_name", "").strip()
    effective_subject_full = subject_full.strip()
    effective_aliases = list(aliases)
    if auto_extract_subject and extracted_subject_full and not effective_subject_full:
        effective_subject_full = extracted_subject_full
    if auto_extract_subject:
        for a in transcript_policy.get("subject_aliases", []):
            if a and a not in effective_aliases:
                effective_aliases.append(a)
    subject_pats = build_subject_patterns(effective_subject_full, subject_email.strip(), effective_aliases)
    private_patterns = compile_patterns(private_patterns_text)
    privileged_keywords = [x.strip() for x in privileged_keywords_text if x.strip()]
    candidates: List[Dict[str, Any]] = []
    seen = set()
    block_meta = []
    if effective_mode == "transcript":
        collect_transcript_policy_candidates(paras=paras, candidates=candidates, seen=seen, subject_pats=subject_pats, transcript_policy=transcript_policy, private_patterns=private_patterns)
        collect_case_document_policy_candidates(paras=paras, subject_pats=subject_pats, transcript_policy=transcript_policy, private_patterns=private_patterns, candidates=candidates, seen=seen)
        if enable_regex:
            for para_idx, p in enumerate(paras):
                collect_regex_candidates_in_para(para_text=p, para_idx=para_idx, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, candidates=candidates, seen=seen, reason_prefix="Transcript policy")
        if enable_ner_backstop and ner_pipe is not None:
            for para_idx, p in enumerate(paras):
                try:
                    persons = find_person_spans(p, ner_pipe)
                except Exception:
                    persons = []
                for s, e, ent_text, label in persons:
                    if mentions_subject(ent_text, subject_pats):
                        continue
                    add_candidate(candidates, seen, para_idx=para_idx, start=s, end=e, text=p, detector="NER", entity_label=label, reason="NER backstop for transcript non-subject person", suggested_decision="review")
    elif effective_mode == "email":
        block_meta = collect_email_chain_candidates(paras=paras, subject_pats=subject_pats, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, privileged_keywords=privileged_keywords, private_patterns=private_patterns, enable_regex=enable_regex, enable_ner_backstop=enable_ner_backstop, ner_pipe=ner_pipe, candidates=candidates, seen=seen)
        collect_case_document_policy_candidates(paras=paras, subject_pats=subject_pats, transcript_policy=transcript_policy, private_patterns=private_patterns, candidates=candidates, seen=seen)
    else:
        collect_generic_candidates(paras=paras, subject_pats=subject_pats, subject_email_lower=subject_email_lower, keep_emails_lower=keep_emails_lower, private_patterns=private_patterns, privileged_keywords=privileged_keywords, enable_regex=enable_regex, enable_ner_backstop=enable_ner_backstop, ner_pipe=ner_pipe, transcript_policy=transcript_policy, candidates=candidates, seen=seen)
    candidates = dedupe_overlaps_prefer_rules(candidates)
    meta = {"detected_mode": detected_mode, "effective_mode": effective_mode, "auto_extracted_subject": extracted_subject_full, "effective_subject_full": effective_subject_full, "effective_aliases": effective_aliases, "keep_emails": sorted(keep_emails_lower), "private_patterns": [p.pattern for p in private_patterns], "privileged_keywords": privileged_keywords, "transcript_policy": {"subject": transcript_policy.get("subject", {}), "participants": transcript_policy.get("participants", []), "third_party_full_names": transcript_policy.get("third_party_full_names", []), "third_party_surnames": transcript_policy.get("third_party_surnames", [])}, "block_meta": block_meta}
    return {"paragraphs": paras, "candidates": candidates, "meta": meta}


# ============================================================
# REVIEW TABLE / OUTPUT HELPERS
# ============================================================
def build_decisions_df(candidates):
    rows = []
    for c in candidates:
        rows.append({"id": c["id"], "decision": c.get("decision", c.get("suggested_decision", "redact")), "detector": c["detector"], "entity_label": c["entity_label"], "text": c["text"], "paragraph": c["para_idx"] + 1, "char_span": f'{c["start"]}:{c["end"]}', "reason": c["reason"]})
    return pd.DataFrame(rows)


def apply_decisions_to_candidates(candidates, edited_df):
    dec_map = {}
    if edited_df is not None and not edited_df.empty:
        for _, r in edited_df.iterrows():
            dec_map[str(r["id"])] = str(r["decision"]).lower().strip()
    out = []
    for c in candidates:
        cc = dict(c)
        cc["decision"] = dec_map.get(c["id"], c.get("decision", "redact"))
        if cc["decision"] not in {"redact", "review", "keep"}:
            cc["decision"] = "review"
        out.append(cc)
    return out


def render_highlighted_html(paragraphs, candidates):
    by_para: Dict[int, List[Dict[str, Any]]] = {}
    for c in candidates:
        by_para.setdefault(c["para_idx"], []).append(c)
    style = '<style>.docwrap {white-space: pre-wrap; font-family: ui-monospace, SFMono-Regular, Menlo, Consolas, monospace; line-height: 1.45;} .hl {padding: 0 2px; border-radius: 3px; border: 1px solid rgba(0,0,0,0.15);} .redact {background: rgba(255, 0, 0, 0.18);} .review {background: rgba(255, 191, 0, 0.22);} .keep {background: rgba(0, 180, 0, 0.18);} .para {margin-bottom: 0.75rem;}</style>'
    parts = [style, '<div class="docwrap">']
    for i, p in enumerate(paragraphs):
        cands = by_para.get(i, [])
        if not cands:
            parts.append(f'<div class="para">{html.escape(p) if p else "&nbsp;"}</div>')
            continue
        priority = {"redact": 0, "review": 1, "keep": 2}
        cands_sorted = sorted(cands, key=lambda c: (c["start"], priority.get(c.get("decision", "review"), 9), -(c["end"] - c["start"])))
        accepted = []
        occupied = []
        for c in cands_sorted:
            s, e = c["start"], c["end"]
            overlaps = any(not (e <= os or s >= oe) for os, oe in occupied)
            if overlaps:
                continue
            accepted.append(c)
            occupied.append((s, e))
        accepted.sort(key=lambda c: c["start"])
        cursor = 0
        para_html = []
        for c in accepted:
            s, e = c["start"], c["end"]
            if cursor < s:
                para_html.append(html.escape(p[cursor:s]))
            cls = c.get("decision", "review")
            title = f'{c["detector"]}:{c["entity_label"]} | {c["reason"]} | {c["id"]}'
            para_html.append(f'<span class="hl {cls}" title="{html.escape(title)}">{html.escape(p[s:e])}</span>')
            cursor = e
        if cursor < len(p):
            para_html.append(html.escape(p[cursor:]))
        parts.append(f'<div class="para">{"".join(para_html) if p else "&nbsp;"}</div>')
    parts.append("</div>")
    return "".join(parts)


def build_redacted_text(paragraphs, candidates):
    by_para: Dict[int, List[Tuple[int, int]]] = {}
    for c in candidates:
        if c.get("decision") == "redact":
            by_para.setdefault(c["para_idx"], []).append((c["start"], c["end"]))
    out_paras = []
    for i, p in enumerate(paragraphs):
        spans = by_para.get(i, [])
        out_paras.append(redact_by_spans(p, spans, token=REDACT_TOKEN))
    return "\n\n".join(out_paras)


def build_redacted_docx_bytes(paragraphs, candidates):
    redacted_text = build_redacted_text(paragraphs, candidates)
    d = Document()
    for p in redacted_text.split("\n\n"):
        d.add_paragraph(p)
    bio = io.BytesIO()
    d.save(bio)
    bio.seek(0)
    return bio.read()


def next_candidate_id(candidates):
    max_n = 0
    for c in candidates:
        m = re.match(r"^C(\d+)$", str(c.get("id", "")))
        if m:
            max_n = max(max_n, int(m.group(1)))
    return f"C{max_n + 1:05d}"


def add_manual_candidate_by_offsets(para_idx, start, end, decision="review"):
    analysis = st.session_state.get("analysis")
    if analysis is None:
        return False
    paragraphs = analysis["paragraphs"]
    if para_idx < 0 or para_idx >= len(paragraphs):
        return False
    para_text = paragraphs[para_idx]
    if start > end:
        start, end = end, start
    start = max(0, min(start, len(para_text)))
    end = max(0, min(end, len(para_text)))
    while start < end and para_text[start].isspace():
        start += 1
    while end > start and para_text[end - 1].isspace():
        end -= 1
    if end <= start:
        return False
    decision = str(decision).lower().strip()
    if decision not in {"redact", "review", "keep"}:
        decision = "review"
    working_candidates = [dict(c) for c in st.session_state.get("working_candidates", analysis["candidates"])]
    for c in working_candidates:
        if c["para_idx"] == para_idx and c["start"] == start and c["end"] == end:
            c["decision"] = decision
            st.session_state["working_candidates"] = working_candidates
            st.session_state["edited_df"] = build_decisions_df(working_candidates)
            st.session_state["final_candidates"] = working_candidates
            st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
            return True
    survivors = [c for c in working_candidates if not (c["para_idx"] == para_idx and overlap((start, end), (c["start"], c["end"])))]
    survivors.append({"id": next_candidate_id(survivors), "para_idx": para_idx, "start": start, "end": end, "text": para_text[start:end], "detector": "MANUAL", "entity_label": "USER_SELECTION", "reason": "User-created manual span", "suggested_decision": decision, "decision": decision})
    survivors = sorted(survivors, key=lambda c: (c["para_idx"], c["start"], c["end"]))
    st.session_state["working_candidates"] = survivors
    st.session_state["edited_df"] = build_decisions_df(survivors)
    st.session_state["final_candidates"] = survivors
    st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
    return True


def add_manual_candidate_by_text(para_idx, selected_text, occurrence=1, decision="review"):
    analysis = st.session_state.get("analysis")
    if analysis is None:
        return False, "No analysis in session."
    paragraphs = analysis["paragraphs"]
    if para_idx < 0 or para_idx >= len(paragraphs):
        return False, "Paragraph index out of range."
    para_text = paragraphs[para_idx]
    needle = selected_text
    if not needle:
        return False, "Enter the exact text you want to mark."
    matches = literal_occurrences(para_text, needle)
    if not matches:
        return False, "That exact text was not found in the chosen paragraph."
    occ = max(1, int(occurrence))
    if occ > len(matches):
        return False, f"Only {len(matches)} matching occurrence(s) found in that paragraph."
    start, end = matches[occ - 1]
    ok = add_manual_candidate_by_offsets(para_idx, start, end, decision=decision)
    return ok, "Manual span added." if ok else "Could not add manual span."


def add_manual_selection_candidate(selection_evt):
    if not selection_evt or selection_evt.get("type") != "create_selection":
        return False
    evt_ts = selection_evt.get("ts")
    if evt_ts == st.session_state.get("last_selection_evt_ts"):
        return False
    st.session_state["last_selection_evt_ts"] = evt_ts
    analysis = st.session_state.get("analysis")
    if analysis is None:
        return False
    paragraphs = analysis["paragraphs"]
    para_idx = int(selection_evt.get("para_idx", -1))
    if para_idx < 0 or para_idx >= len(paragraphs):
        return False
    para_text = paragraphs[para_idx]
    start = int(selection_evt.get("start", -1))
    end = int(selection_evt.get("end", -1))
    if start > end:
        start, end = end, start
    start = max(0, min(start, len(para_text)))
    end = max(0, min(end, len(para_text)))
    while start < end and para_text[start].isspace():
        start += 1
    while end > start and para_text[end - 1].isspace():
        end -= 1
    if end <= start:
        return False
    decision = str(selection_evt.get("decision", "review")).lower().strip()
    if decision not in {"redact", "review"}:
        decision = "review"
    working_candidates = [dict(c) for c in st.session_state.get("working_candidates", analysis["candidates"])]
    for c in working_candidates:
        if c["para_idx"] == para_idx and c["start"] == start and c["end"] == end:
            c["decision"] = decision
            st.session_state["working_candidates"] = working_candidates
            st.session_state["edited_df"] = build_decisions_df(working_candidates)
            st.session_state["final_candidates"] = working_candidates
            st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
            return True
    survivors = [c for c in working_candidates if not (c["para_idx"] == para_idx and overlap((start, end), (c["start"], c["end"])))]
    survivors.append({"id": next_candidate_id(survivors), "para_idx": para_idx, "start": start, "end": end, "text": para_text[start:end], "detector": "MANUAL", "entity_label": "USER_SELECTION", "reason": "User-selected span from interactive preview", "suggested_decision": decision, "decision": decision})
    survivors = sorted(survivors, key=lambda c: (c["para_idx"], c["start"], c["end"]))
    st.session_state["working_candidates"] = survivors
    st.session_state["edited_df"] = build_decisions_df(survivors)
    st.session_state["final_candidates"] = survivors
    st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
    return True


def sync_programmatic_decision_change(click_evt):
    if not click_evt:
        return False
    evt_id = click_evt.get("id")
    evt_ts = click_evt.get("ts")
    evt_decision = str(click_evt.get("decision", "")).strip().lower()
    if not evt_id or evt_decision not in {"redact", "review", "keep"}:
        return False
    if evt_ts == st.session_state.get("last_click_evt_ts"):
        return False
    st.session_state["last_click_evt_ts"] = evt_ts
    edited_df = st.session_state.get("edited_df")
    working_candidates = st.session_state.get("working_candidates")
    if edited_df is None or working_candidates is None:
        return False
    edited_df = edited_df.copy()
    mask = edited_df["id"] == evt_id
    if not mask.any():
        return False
    edited_df.loc[mask, "decision"] = evt_decision
    st.session_state["edited_df"] = edited_df
    final_candidates = apply_decisions_to_candidates(working_candidates, edited_df)
    st.session_state["working_candidates"] = final_candidates
    st.session_state["final_candidates"] = final_candidates
    st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
    return True


def sync_preview_event(evt):
    if not evt:
        return False
    evt_type = evt.get("type", "update_decision")
    if evt_type == "create_selection":
        return add_manual_selection_candidate(evt)
    return sync_programmatic_decision_change(evt)


# ============================================================
# STREAMLIT UI
# ============================================================
st.set_page_config(page_title="DSAR Redaction Review", layout="wide")
st.title("DSAR Redaction Review")
st.caption("Upload a Word document, review highlighted PII, confirm redactions, and download.")

if "editor_version" not in st.session_state:
    st.session_state["editor_version"] = 0

with st.sidebar:
    st.header("Configuration")

    document_mode = st.selectbox("Document mode", options=["auto", "email", "transcript", "generic"], index=0, help="Auto tries to detect transcript vs email chain structure.")
    auto_extract_subject = st.checkbox("Auto-detect data subject from transcript 'Employee:' line", value=True)
    subject_full = st.text_input("Data Subject Full Name", value=DEFAULT_DATA_SUBJECT_FULL)
    subject_email = st.text_input("Data Subject Email", value=DEFAULT_DATA_SUBJECT_EMAIL)
    aliases_text = st.text_area("Aliases (one per line)", value="\n".join(DEFAULT_ALIASES), height=110)
    keep_emails_text = st.text_area("Allowlisted emails (one per line)", value="\n".join(sorted(DEFAULT_KEEP_EMAILS)), height=110)

    # Patterns and keywords loaded from config — not exposed to end users
    private_patterns_text = "\n".join(DEFAULT_PRIVATE_PATTERNS)
    privileged_keywords_text = "\n".join(DEFAULT_PRIVILEGED_KEYWORDS)

    enable_regex = st.checkbox("Use regex for emails / phones", value=True)
    enable_ner_backstop = st.checkbox("Use NER backstop for extra names", value=True)

uploaded = st.file_uploader("Upload source DOCX", type=["docx"])

col_a, col_b = st.columns([1, 1])
with col_a:
    analyze_btn = st.button("Analyse document", type="primary", disabled=uploaded is None)
with col_b:
    clear_btn = st.button("Clear session")

if clear_btn:
    for k in ["analysis", "edited_df", "final_candidates", "working_candidates", "last_click_evt_ts", "last_selection_evt_ts", "candidate_editor_state", "clickable_preview_value"]:
        st.session_state.pop(k, None)
    st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
    st.rerun()

if analyze_btn and uploaded is not None:
    aliases = [x.strip() for x in aliases_text.splitlines() if x.strip()]
    keep_emails = {x.strip() for x in keep_emails_text.splitlines() if x.strip()}
    private_patterns_lines = [x.strip() for x in private_patterns_text.splitlines() if x.strip()]
    privileged_keywords_lines = [x.strip() for x in privileged_keywords_text.splitlines() if x.strip()]

    with st.spinner("Loading models..."):
        ner_pipe = load_ner_pipeline() if enable_ner_backstop else None

    with st.spinner("Reading DOCX and generating candidate redactions..."):
        doc = Document(uploaded)
        paras = [p.text for p in doc.paragraphs if p.text is not None]
        analysis = analyze_docx_paragraphs(paras, subject_full=subject_full, subject_email=subject_email, aliases=aliases, keep_emails=keep_emails, private_patterns_text=private_patterns_lines, privileged_keywords_text=privileged_keywords_lines, enable_regex=enable_regex, enable_ner_backstop=enable_ner_backstop, auto_extract_subject=auto_extract_subject, document_mode=document_mode, ner_pipe=ner_pipe)
        st.session_state["analysis"] = analysis
        st.session_state["working_candidates"] = [dict(c) for c in analysis["candidates"]]
        st.session_state["edited_df"] = build_decisions_df(st.session_state["working_candidates"])
        st.session_state["final_candidates"] = [dict(c) for c in st.session_state["working_candidates"]]
        st.session_state["editor_version"] = st.session_state.get("editor_version", 0) + 1
        st.rerun()

if "analysis" in st.session_state:
    analysis = st.session_state["analysis"]
    paragraphs = analysis["paragraphs"]
    candidates = analysis["candidates"]

    st.subheader("Candidate redactions")
    st.write(f"Found **{len(candidates)}** candidate redactions.")

    if len(candidates) == 0:
        st.info("No candidate redactions were generated with current settings.")
    else:
        working_candidates = st.session_state.get("working_candidates", candidates)
        current_candidates = st.session_state.get("final_candidates", working_candidates)

        st.subheader("Highlighted preview")

        if clickable_preview_available():
            preview_key = f"clickable_preview_{st.session_state.get('editor_version', 0)}"
            click_evt = clickable_redaction_preview(paragraphs=paragraphs, candidates=current_candidates, key=preview_key, default={})
            if click_evt and sync_preview_event(click_evt):
                st.rerun()
            st.caption("Click an existing highlighted span to change its classification, or drag-select any text in the preview to create a new manual Redact/Review span.")
        else:
            preview_html = render_highlighted_html(paragraphs, current_candidates)
            st.components.v1.html(preview_html, height=520, scrolling=True)

        with st.expander("Manual span fallback (use this if the interactive preview does not load)", expanded=not clickable_preview_available()):
            st.caption("Use this fallback to add a manual span by exact text match or by character offsets.")
            para_options = [f"{i+1}: {normalize(p)[:120]}" for i, p in enumerate(paragraphs)]
            chosen_para_label = st.selectbox("Paragraph", options=para_options, key="manual_para_select")
            chosen_para_idx = para_options.index(chosen_para_label)
            chosen_para_text = paragraphs[chosen_para_idx]
            st.text_area("Chosen paragraph text", chosen_para_text, height=180, disabled=True, key="manual_para_text_preview")
            manual_mode = st.radio("Add span by", options=["exact text", "character offsets"], horizontal=True, key="manual_add_mode")
            manual_decision = st.selectbox("Classification", options=["redact", "review"], key="manual_decision_select")
            if manual_mode == "exact text":
                manual_text = st.text_input("Exact text to mark", key="manual_exact_text", help="Paste the exact text exactly as it appears in the paragraph.")
                match_count = len(literal_occurrences(chosen_para_text, manual_text)) if manual_text else 0
                st.caption(f"Matches in this paragraph: {match_count}")
                occurrence = st.number_input("Occurrence number", min_value=1, value=1, step=1, key="manual_occurrence_number")
                if st.button("Add manual span from text", key="manual_add_from_text_btn"):
                    ok, msg = add_manual_candidate_by_text(chosen_para_idx, manual_text, occurrence=int(occurrence), decision=manual_decision)
                    if ok:
                        st.success(msg)
                        st.rerun()
                    else:
                        st.error(msg)
            else:
                cstart, cend = st.columns(2)
                with cstart:
                    manual_start = st.number_input("Start offset", min_value=0, max_value=max(0, len(chosen_para_text)), value=0, step=1, key="manual_start_offset")
                with cend:
                    manual_end = st.number_input("End offset", min_value=0, max_value=max(0, len(chosen_para_text)), value=min(10, len(chosen_para_text)), step=1, key="manual_end_offset")
                start_i = int(manual_start)
                end_i = int(manual_end)
                lo, hi = sorted((start_i, end_i))
                st.text_input("Selected snippet preview", value=chosen_para_text[lo:hi], disabled=True, key="manual_selected_snippet_preview")
                if st.button("Add manual span from offsets", key="manual_add_from_offsets_btn"):
                    ok = add_manual_candidate_by_offsets(chosen_para_idx, lo, hi, decision=manual_decision)
                    if ok:
                        st.success("Manual span added.")
                        st.rerun()
                    else:
                        st.error("Could not add manual span for those offsets.")

        editor_key = f"candidate_editor_{st.session_state.get('editor_version', 0)}"
        df = st.session_state.get("edited_df", build_decisions_df(working_candidates))
        edited_df = st.data_editor(df, num_rows="fixed", use_container_width=True, hide_index=True, column_config={"decision": st.column_config.SelectboxColumn("decision", options=["redact", "review", "keep"], required=True, help="redact = apply redaction, review = flag only, keep = no redaction"), "id": st.column_config.TextColumn("id", disabled=True, width="small"), "detector": st.column_config.TextColumn("detector", disabled=True, width="small"), "entity_label": st.column_config.TextColumn("entity_label", disabled=True, width="small"), "text": st.column_config.TextColumn("text", disabled=True, width="large"), "paragraph": st.column_config.NumberColumn("paragraph", disabled=True, width="small"), "char_span": st.column_config.TextColumn("char_span", disabled=True, width="small"), "reason": st.column_config.TextColumn("reason", disabled=True, width="large")}, key=editor_key)

        final_candidates = apply_decisions_to_candidates(working_candidates, edited_df)
        st.session_state["edited_df"] = edited_df
        st.session_state["working_candidates"] = final_candidates
        st.session_state["final_candidates"] = final_candidates

        c1, c2, c3 = st.columns(3)
        with c1:
            st.metric("Redact", sum(1 for c in final_candidates if c["decision"] == "redact"))
        with c2:
            st.metric("Review", sum(1 for c in final_candidates if c["decision"] == "review"))
        with c3:
            st.metric("Keep", sum(1 for c in final_candidates if c["decision"] == "keep"))

        st.subheader("Redacted text output")
        redacted_text = build_redacted_text(paragraphs, final_candidates)
        st.text_area("Final redacted text", redacted_text, height=320)

        audit_df = pd.DataFrame(final_candidates)
        audit_csv = audit_df.to_csv(index=False).encode("utf-8")
        redacted_docx_bytes = build_redacted_docx_bytes(paragraphs, final_candidates)

        d1, d2, d3 = st.columns(3)
        with d1:
            st.download_button("Download redacted .txt", data=redacted_text.encode("utf-8"), file_name="dsar_redacted_output.txt", mime="text/plain")
        with d2:
            st.download_button("Download audit .csv", data=audit_csv, file_name="dsar_redaction_audit.csv", mime="text/csv")
        with d3:
            st.download_button("Download redacted .docx", data=redacted_docx_bytes, file_name="dsar_redacted_output.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with st.expander("Analysis metadata / detected policy"):
            st.json(analysis.get("meta", {}))
else:
    st.info("Upload a DOCX and click **Analyse document** to start.")